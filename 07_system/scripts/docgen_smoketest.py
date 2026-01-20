from docx import Document
from reportlab.pdfgen import canvas

docx_path = "/Users/olivermarroquin/secondbrain/00_inbox/smoketest.docx"
pdf_path  = "/Users/olivermarroquin/secondbrain/00_inbox/smoketest.pdf"

# DOCX
doc = Document()
doc.add_heading("Docgen Smoketest", level=1)
doc.add_paragraph("If you can read this, python-docx works.")
doc.save(docx_path)

# PDF
c = canvas.Canvas(pdf_path)
c.setTitle("Docgen Smoketest")
c.drawString(72, 720,"If you can read this, reportlab works.")
c.save()

print("Wrote:")
print(docx_path)
print(pdf_path)

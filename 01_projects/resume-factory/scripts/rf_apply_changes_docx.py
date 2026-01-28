#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

from docx import Document  # type: ignore


def die(msg: str, code: int = 2) -> None:
    print(f"ERROR: {msg}", file=sys.stderr)
    sys.exit(code)


def read_json(p: Path) -> dict:
    try:
        return json.loads(p.read_text(encoding="utf-8"))
    except FileNotFoundError:
        die(f"Missing required file: {p}")
    except json.JSONDecodeError as e:
        die(f"Invalid JSON: {p} ({e})")

def norm_heading(s: str) -> str:
    return (s or "").strip().rstrip(":").upper()


def heading_matches(doc_heading: str, target: str) -> bool:
    # Deterministic matching:
    # - exact match after normalization
    # - doc heading endswith target (e.g., "PROFESSIONAL SUMMARY" endswith "SUMMARY")
    # - doc heading contains target as a whole word-ish token
    dh = norm_heading(doc_heading)
    tg = norm_heading(target)
    if not dh or not tg:
        return False
    if dh == tg:
        return True
    if dh.endswith(tg):
        return True
    # token containment (space-delimited)
    return f" {tg} " in f" {dh} "


def apply_add(doc: Document, section: str, text: str) -> None:
    for p in doc.paragraphs:
        if heading_matches(p.text, section):
            doc.add_paragraph(text)
            return
    die(f"Section not found for ADD: {section}")


def apply_replace_section(doc: Document, section: str, paras: list[str]) -> None:
    in_section = False
    start_idx = None
    end_idx = None

    for i, p in enumerate(doc.paragraphs):
        if heading_matches(p.text, section):
            in_section = True
            start_idx = i + 1
            continue
        if in_section and p.text.strip().endswith(":"):
            end_idx = i
            break

    if start_idx is None:
        die(f"Section not found for REPLACE_SECTION: {section}")

    if end_idx is None:
        end_idx = len(doc.paragraphs)

    for _ in range(end_idx - start_idx):
        doc.paragraphs[start_idx]._element.getparent().remove(doc.paragraphs[start_idx]._element)

    for t in paras:
        doc.add_paragraph(t)


def main() -> None:
    ap = argparse.ArgumentParser(prog="rf_apply_changes_docx")
    ap.add_argument("--app", required=True, help="Job folder path (APP)")
    ap.add_argument("--root", default="~/secondbrain")
    ap.add_argument("--force", action="store_true")
    args = ap.parse_args()

    app = Path(args.app).expanduser().resolve()
    pipe = app / "resume_refs" / "resume_pipeline"

    patches_p = pipe / "patches.json"
    if not patches_p.exists():
        die(f"Missing patches.json: {patches_p}")

    patches = read_json(patches_p)
    if patches.get("schema") != "rf_patches_v1":
        die(f"Invalid patches schema: {patches.get('schema')}")

    template_dir = None
    job_json = read_json(pipe / "job.json")
    sel_json = read_json(pipe / "selection.json")
    try:
        template_dir = Path(sel_json["selected"]["template_dir"])
    except Exception:
        die("selection.json missing selected.template_dir")

    master = template_dir / "resume-master.docx"
    if not master.exists():
        die(f"Missing resume-master.docx: {master}")

    out_dir = Path("~/secondbrain/05_outputs").expanduser() / "resumes"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_doc = out_dir / "resume.docx"

    doc = Document(str(master))

    for item in patches.get("patches", []):
        op = item.get("op")
        section = item.get("section")

        if op == "ADD":
            apply_add(doc, section, item["to_text"])
        elif op == "REPLACE_SECTION":
            paras = item.get("to_paragraphs")
            if not isinstance(paras, list):
                die(f"REPLACE_SECTION missing to_paragraphs for section {section}")
            apply_replace_section(doc, section, paras)
        else:
            die(f"Unsupported op in apply: {op}")

    doc.save(str(out_doc))
    print(str(out_doc))


if __name__ == "__main__":
    main()
